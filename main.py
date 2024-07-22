import os
import argparse
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time
from Content_Creation import ContentCreation
from Image import ImageHandler
import io
import ast

class Keyword:
    def __init__(self, word, avg_monthly_searches, competition):
        self.word = word
        self.avg_monthly_searches = avg_monthly_searches
        self.competition = competition

    def __repr__(self):
        return f"Keyword(word='{self.word}', avg_monthly_searches={self.avg_monthly_searches}, competition='{self.competition}')"

def parse_keywords(keywords_string):
    try:
        keywords_list = ast.literal_eval(keywords_string)
        return [Keyword(k['word'], k['avg_monthly_searches'], k['competition']) for k in keywords_list]
    except (SyntaxError, ValueError) as e:
        print(f"Error parsing keywords: {e}")
        print("Please ensure the keyword string is properly formatted.")
        return None

def create_word_document(content, images, meta_description):
    doc = Document()
    
    # Add title
    title = content.split('\n')[0].strip('#').strip()
    doc.add_heading(title, 0)

    # Add meta description
    if meta_description:
        doc.add_paragraph(f"Meta Description: {meta_description}")

    # Add banner image
    if images['banner']:
        image_stream = io.BytesIO(images['banner'][0]['data'])
        doc.add_picture(image_stream, width=Inches(6))
        caption = doc.add_paragraph(f"Image by {images['banner'][0]['photographer']}")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("Banner image not available")

    # Process content
    body_image_index = 0
    for line in content.split('\n')[1:]:  # Skip the title
        if line.startswith('#'):
            level = line.count('#')
            text = line.strip('#').strip()
            doc.add_heading(text, level)
        elif line.strip() == '!IMAGE!' and images['body'] and body_image_index < len(images['body']):
            img = images['body'][body_image_index]
            image_stream = io.BytesIO(img['data'])
            doc.add_picture(image_stream, width=Inches(4))
            caption = doc.add_paragraph(f"Image by {img['photographer']}")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            body_image_index += 1
        else:
            doc.add_paragraph(line)

    return doc

def generate_blog_post(topic, keywords, purpose, section_count, words_per_section, update_progress):
    update_progress(0, "Initializing blog generation...")
    
    start_time = time.time()

    keywords = parse_keywords(keywords)
    if not keywords:
        update_progress(100, "Failed to parse keywords. Exiting.")
        return None

    content_creator = ContentCreation()
    image_handler = ImageHandler()

    update_progress(10, f"Generating blog content for topic: {topic}")
    content = content_creator.create_content(topic, keywords, section_count, words_per_section, purpose)

    if not content:
        update_progress(100, "Failed to generate content. Exiting.")
        return None

    update_progress(50, "Content generated. Extracting meta description...")
    meta_description = content_creator.extract_meta_description(content)
    content = content_creator.remove_meta_description(content)

    update_progress(60, "Downloading images...")
    images = {
        'banner': image_handler.search_and_download(topic, image_type='banner'),
        'body': image_handler.search_and_download(topic, image_type='body', max_results=3)
    }

    update_progress(80, "Creating Word document...")
    doc = create_word_document(content, images, meta_description)
    
    filename = f"{topic.replace(' ', '_')}_blog.docx"
    file_path = os.path.join(os.getcwd(), filename)
    doc.save(file_path)

    end_time = time.time()
    update_progress(100, f"Blog post generated and saved as '{filename}'")
    print(f"Total runtime: {end_time - start_time:.2f} seconds")
    
    return file_path

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate an SEO-optimized blog post")
    parser.add_argument("topic", help="The main topic of the blog post")
    parser.add_argument("keywords", help="JSON string of keywords with their stats")
    parser.add_argument("purpose", help="The purpose of the article/blog")
    parser.add_argument("--section-count", type=int, default=5, help="Number of main sections in the blog")
    parser.add_argument("--words-per-section", type=int, default=300, help="Approximate words per section")
    args = parser.parse_args()

    def dummy_update_progress(progress, message):
        print(f"Progress: {progress}% - {message}")

    generate_blog_post(args.topic, args.keywords, args.purpose, args.section_count, args.words_per_section, dummy_update_progress)